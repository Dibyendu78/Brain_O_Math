import csv
import io
import json
import docx

from django.contrib.auth import authenticate
from django.core.mail import send_mail
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt

from Account.decorators import jwt_required
from Account.models import User
from Account.views import issue_login_response, send_results_published_email
from Corrdinator.models import CoordinatorProfile
from Corrdinator.views import _profile_dict, _student_dict
from Registartion.models import RegistrationPayment, RegistrationSettings, Student, generate_roll_number



@jwt_required(roles=[User.ADMIN])
def admin_dashboard(request):
    return render(request, "public/admin.html")


def _body(request):
    if request.body:
        try:
            return json.loads(request.body.decode("utf-8"))
        except json.JSONDecodeError:
            return {}
    return request.POST.dict()


def _payment_dict(payment):
    profile = payment.coordinator
    students = list(profile.students.order_by("created_at"))
    return {
        "_id": payment.id,
        "registrationId": payment.registration_id,
        "status": payment.status,
        "paymentStatus": payment.status,
        "utr": payment.utr,
        "venue": payment.venue or "Doon Heritage School, Siliguri",
        "totalAmount": payment.total_amount,
        "school": _profile_dict(profile),
        "students": [_student_dict(student) for student in students],
        "createdAt": payment.submitted_at.isoformat() if payment.submitted_at else "",
    }


@csrf_exempt
def api_admin_login(request):
    data = _body(request)
    username = data.get("username", "").strip()
    password = data.get("password", "")
    user = authenticate(request, username=username, password=password)
    if not user:
        try:
            email_user = User.objects.get(username=username)
            user = authenticate(request, username=email_user.email, password=password)
        except User.DoesNotExist:
            user = None
    if not user or (user.role != User.ADMIN and not user.is_superuser):
        return JsonResponse({"success": False, "message": "Invalid admin credentials"}, status=401)
    if not user.is_active:
        return JsonResponse({"success": False, "message": "Admin account is inactive"}, status=403)
    return issue_login_response(request, user, "/admin-control/")


@jwt_required(roles=[User.ADMIN])
def api_stats(request):
    payments = RegistrationPayment.objects.all()
    return JsonResponse(
        {
            "success": True,
            "data": {
                "totalRegistrations": payments.count(),
                "totalStudents": Student.objects.count(),
                "totalRevenue": sum(payment.total_amount for payment in payments.filter(status="verified")),
            },
        }
    )


@jwt_required(roles=[User.ADMIN])
def api_registrations(request):
    status = request.GET.get("status")
    venue = request.GET.get("venue")
    payments = RegistrationPayment.objects.select_related("coordinator", "coordinator__user").order_by("-updated_at")
    if status:
        payments = payments.filter(status=status)
    if venue:
        payments = payments.filter(venue=venue)
    return JsonResponse({"success": True, "data": [_payment_dict(payment) for payment in payments], "pagination": {"total": payments.count(), "page": 1}})


@csrf_exempt
@jwt_required(roles=[User.ADMIN])
def api_registration_status_update(request, registration_id):
    data = _body(request)
    status = data.get("status")
    if status not in {"pending", "submitted", "verified", "rejected"}:
        return JsonResponse({"success": False, "message": "Invalid status"}, status=400)
    try:
        payment = RegistrationPayment.objects.get(registration_id=registration_id)
    except RegistrationPayment.DoesNotExist:
        return JsonResponse({"success": False, "message": "Registration not found"}, status=404)
    payment.status = status
    payment.save(update_fields=["status", "updated_at"])
    return JsonResponse({"success": True, "data": _payment_dict(payment)})


@jwt_required(roles=[User.ADMIN])
def api_students(request):
    from django.db.models import Q
    qs = Student.objects.select_related("coordinator", "coordinator__user", "coordinator__payment").order_by("student_class", "name")
    cls = request.GET.get("class")
    venue = request.GET.get("venue")
    subject = request.GET.get("subject")
    status = request.GET.get("status")
    school = request.GET.get("school") or request.GET.get("school_id")
    search = request.GET.get("search")
    marks_status = request.GET.get("marks_status")

    if cls:
        qs = qs.filter(student_class=str(cls).strip())
    if venue:
        qs = qs.filter(venue__iexact=venue.strip())
    if subject:
        qs = qs.filter(subjects__icontains=subject.strip())
    if status:
        qs = qs.filter(coordinator__payment__status=status.strip())
    if school:
        school_str = str(school).strip()
        if school_str.isdigit():
            qs = qs.filter(coordinator_id=int(school_str))
        else:
            qs = qs.filter(coordinator__school_name__icontains=school_str)
    if search:
        search_str = search.strip()
        qs = qs.filter(
            Q(name__icontains=search_str) |
            Q(roll_number__icontains=search_str) |
            Q(student_id__icontains=search_str) |
            Q(coordinator__school_name__icontains=search_str)
        )
    if marks_status in ("completed", "ready"):
        qs = qs.filter(
            Q(english_marks__isnull=False) |
            Q(math_marks__isnull=False) |
            Q(science_marks__isnull=False) |
            Q(cs_marks__isnull=False)
        )
    elif marks_status == "pending":
        qs = qs.filter(
            english_marks__isnull=True,
            math_marks__isnull=True,
            science_marks__isnull=True,
            cs_marks__isnull=True
        )

    data = []
    for student in qs:
        item = _student_dict(student)
        item["school"] = _profile_dict(student.coordinator)
        if hasattr(student.coordinator, "payment"):
            item["paymentStatus"] = student.coordinator.payment.status
            item["status"] = student.coordinator.payment.status
        data.append(item)
    return JsonResponse({"success": True, "data": data, "pagination": {"total": qs.count(), "page": 1}})



@csrf_exempt
@jwt_required(roles=[User.ADMIN])
def api_marks(request, student_id=None):
    data = _body(request)
    if student_id:
        students = [
            {
                "studentId": student_id,
                "englishMarks": data.get("englishMarks"),
                "mathMarks": data.get("mathMarks"),
                "scienceMarks": data.get("scienceMarks"),
                "csMarks": data.get("csMarks"),
            }
        ]
    else:
        students = data.get("students", [])

    success_count = 0
    for row in students:
        try:
            student = Student.objects.get(id=row.get("studentId"))
        except Student.DoesNotExist:
            continue
        student.english_marks = _optional_int(row.get("englishMarks"))
        student.math_marks = _optional_int(row.get("mathMarks"))
        student.science_marks = _optional_int(row.get("scienceMarks"))
        student.cs_marks = _optional_int(row.get("csMarks"))
        student.save(update_fields=["english_marks", "math_marks", "science_marks", "cs_marks", "updated_at"])
        success_count += 1
    return JsonResponse(
        {
            "success": True,
            "data": {
                "successCount": success_count,
                "total": len(students),
                "failureCount": len(students) - success_count,
            },
        }
    )


@csrf_exempt
@jwt_required(roles=[User.ADMIN])
def api_edit_student(request, student_id=None):
    data = _body(request)
    rows = data.get("students", [{"studentId": student_id, "name": data.get("name")}])
    success_count = 0
    for row in rows:
        try:
            student = Student.objects.get(id=row.get("studentId"))
        except Student.DoesNotExist:
            continue
        student.name = (row.get("name") or student.name).strip()
        student.save(update_fields=["name", "updated_at"])
        success_count += 1
    return JsonResponse({"success": True, "data": {"successCount": success_count, "total": len(rows), "failureCount": len(rows) - success_count}})


@csrf_exempt
@jwt_required(roles=[User.ADMIN])
def api_registration_open(request):
    settings = RegistrationSettings.current()
    if request.method in {"POST", "PUT"}:
        data = _body(request)
        if "status" in data:
            settings.is_open = data["status"] == "open"
        if "isOpen" in data:
            settings.is_open = bool(data["isOpen"])
        if "message" in data:
            settings.message = data["message"]
        settings.save()
    return JsonResponse(
        {
            "success": True,
            "data": {
                "isOpen": settings.is_open,
                "status": "open" if settings.is_open else "closed",
                "message": settings.message,
                "resultsPublished": settings.results_published,
            },
        }
    )


import threading

@csrf_exempt
@jwt_required(roles=[User.ADMIN])
def api_publish_results(request):
    settings = RegistrationSettings.current()
    was_published = settings.results_published
    settings.results_published = True
    settings.save(update_fields=["results_published", "updated_at"])
    if not was_published:
        def _send_all_emails():
            for profile in CoordinatorProfile.objects.select_related("user").all():
                if profile.user and profile.user.email:
                    try:
                        send_results_published_email(request, profile)
                    except Exception:
                        pass
        threading.Thread(target=_send_all_emails, daemon=True).start()
    return JsonResponse({"success": True, "message": "Results published successfully.", "data": {"resultsPublished": settings.results_published}})




@csrf_exempt
@jwt_required(roles=[User.ADMIN])
def api_unpublish_results(request):
    settings = RegistrationSettings.current()
    settings.results_published = False
    settings.save(update_fields=["results_published", "updated_at"])
    return JsonResponse({"success": True, "message": "Results hidden successfully.", "data": {"resultsPublished": settings.results_published}})


@jwt_required(roles=[User.ADMIN])
def api_result_stats(request):
    from django.db.models import Q
    total = Student.objects.count()
    with_marks = Student.objects.filter(
        Q(english_marks__isnull=False) |
        Q(math_marks__isnull=False) |
        Q(science_marks__isnull=False) |
        Q(cs_marks__isnull=False)
    ).count()
    pending = total - with_marks
    percentage = round((with_marks / total * 100), 1) if total > 0 else 0
    return JsonResponse({
        "success": True,
        "data": {
            "summary": {
                "totalStudents": total,
                "withMarks": with_marks,
                "pending": pending,
                "percentage": percentage,
            }
        }
    })


@jwt_required(roles=[User.ADMIN])
def api_admit_card_stats(request):
    total = Student.objects.count()
    released = Student.objects.filter(admit_card_released=True).count()
    pending = total - released
    percentage = round((released / total * 100), 1) if total > 0 else 0
    return JsonResponse({
        "success": True,
        "data": {
            "summary": {
                "totalStudents": total,
                "released": released,
                "pending": pending,
                "percentage": percentage,
            }
        }
    })


@csrf_exempt
@jwt_required(roles=[User.ADMIN])
def api_release_admit_card(request, student_id=None, school_id=None):
    if student_id:
        qs = Student.objects.filter(id=student_id)
    elif school_id:
        qs = Student.objects.filter(coordinator_id=school_id)
    else:
        qs = Student.objects.all()
    released = 0
    single_roll = ""
    for student in qs:
        # Always normalize the roll number when releasing, including records
        # that were generated with the old BOM-4-0009 format.
        student.roll_number = generate_roll_number(student.student_class, student.id)
        student.admit_card_released = "revoke" not in request.path
        student.save(update_fields=["roll_number", "admit_card_released", "updated_at"])
        released += 1
        single_roll = student.roll_number
    return JsonResponse({"success": True, "data": {"released": released, "rollNumber": single_roll}})


@jwt_required(roles=[User.ADMIN])
def api_export(request):
    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = 'attachment; filename="brain-o-math-export.csv"'
    writer = csv.writer(response)
    writer.writerow(["Registration ID", "School", "Coordinator", "Email", "Students", "Venue", "Amount", "Status", "UTR"])
    for payment in RegistrationPayment.objects.select_related("coordinator", "coordinator__user"):
        writer.writerow([
            payment.registration_id,
            payment.coordinator.school_name,
            payment.coordinator.coordinator_name,
            payment.coordinator.user.email,
            payment.coordinator.students.count(),
            payment.venue or "Doon Heritage School, Siliguri",
            payment.total_amount,
            payment.status,
            payment.utr,
        ])
    return response


@jwt_required(roles=[User.ADMIN])
def api_export_students(request):
    qs = Student.objects.select_related("coordinator", "coordinator__user", "coordinator__payment").order_by("student_class", "name")
    cls = request.GET.get("class")
    venue = request.GET.get("venue")
    subject = request.GET.get("subject")
    status = request.GET.get("status")

    if cls:
        qs = qs.filter(student_class=cls)
    if venue:
        qs = qs.filter(venue=venue)
    if subject:
        qs = qs.filter(subjects__icontains=subject)
    if status:
        qs = qs.filter(coordinator__payment__status=status)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = 'attachment; filename="students-export.csv"'
    writer = csv.writer(response)
    writer.writerow([
        "Student ID", "Roll Number", "Name", "Class", "Category", "Subjects", 
        "Venue", "Fee", "Parent Name", "Parent Contact", "School Name", 
        "Coordinator Name", "Coordinator Email", "Payment Status", "Admit Card Released"
    ])

    for s in qs:
        payment_status = s.coordinator.payment.status if hasattr(s.coordinator, "payment") else "pending"
        writer.writerow([
            s.student_id,
            s.roll_number or "N/A",
            s.name,
            s.student_class,
            s.category,
            s.subjects,
            s.venue or "Doon Heritage School, Siliguri",
            s.fee,
            s.parent_name or "N/A",
            s.parent_contact or "N/A",
            s.coordinator.school_name,
            s.coordinator.coordinator_name,
            s.coordinator.user.email,
            payment_status,
            "Yes" if s.admit_card_released else "No",
        ])
    return response


@jwt_required(roles=[User.ADMIN])
def api_export_docx(request):
    doc = docx.Document()
    doc.add_heading('Brain-O-Math Olympiad 2026 - Registration & Venue Report', 0)

    venue_filter = request.GET.get("venue")
    payments = RegistrationPayment.objects.select_related("coordinator", "coordinator__user").order_by("-updated_at")
    if venue_filter:
        payments = payments.filter(venue=venue_filter)

    doc.add_heading('Registration Summaries', level=1)

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Reg ID'
    hdr_cells[1].text = 'School Name'
    hdr_cells[2].text = 'Coordinator'
    hdr_cells[3].text = 'Venue'
    hdr_cells[4].text = 'Students'
    hdr_cells[5].text = 'Amount (₹)'
    hdr_cells[6].text = 'Status'
    hdr_cells[7].text = 'UTR'

    for payment in payments:
        row_cells = table.add_row().cells
        row_cells[0].text = str(payment.registration_id)
        row_cells[1].text = str(payment.coordinator.school_name)
        row_cells[2].text = str(payment.coordinator.coordinator_name)
        row_cells[3].text = str(payment.venue or "Doon Heritage School, Siliguri")
        row_cells[4].text = str(payment.coordinator.students.count())
        row_cells[5].text = str(payment.total_amount)
        row_cells[6].text = str(payment.status)
        row_cells[7].text = str(payment.utr or "N/A")

    doc.add_paragraph('')
    doc.add_heading('Candidate Details', level=1)

    students = Student.objects.select_related("coordinator").order_by("student_class", "name")
    if venue_filter:
        students = students.filter(venue=venue_filter)

    st_table = doc.add_table(rows=1, cols=8)
    st_table.style = 'Table Grid'
    st_hdr = st_table.rows[0].cells
    st_hdr[0].text = 'Student ID'
    st_hdr[1].text = 'Name'
    st_hdr[2].text = 'Class'
    st_hdr[3].text = 'Category'
    st_hdr[4].text = 'Subjects'
    st_hdr[5].text = 'Chosen Venue'
    st_hdr[6].text = 'School'
    st_hdr[7].text = 'Roll No'

    for st in students:
        row = st_table.add_row().cells
        row[0].text = str(st.student_id)
        row[1].text = str(st.name)
        row[2].text = str(st.student_class)
        row[3].text = str(st.category)
        row[4].text = str(st.subjects)
        row[5].text = str(st.venue or "Doon Heritage School, Siliguri")
        row[6].text = str(st.coordinator.school_name)
        row[7].text = str(st.roll_number or "N/A")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="brain-o-math-registrations-venue-report.docx"'
    return response


def _optional_int(value):
    if value in ("", None):
        return None
    return int(value)


@jwt_required(roles=[User.ADMIN])
def api_schools(request):
    seen = set()
    schools = []
    for profile in CoordinatorProfile.objects.select_related("user").order_by("school_name"):
        sname = (profile.school_name or "").strip()
        if sname and sname.lower() not in seen:
            seen.add(sname.lower())
            item = _profile_dict(profile)
            item["schoolName"] = sname
            schools.append(item)
    return JsonResponse({"success": True, "data": schools})


@csrf_exempt
@jwt_required(roles=[User.ADMIN])
def api_send_message_to_coordinators(request):
    data = _body(request)
    message_content = data.get("message", "").strip()
    
    if not message_content:
        return JsonResponse({"success": False, "message": "Message cannot be empty"}, status=400)
    
    coordinators = CoordinatorProfile.objects.select_related("user").all()
    success_count = 0
    
    for profile in coordinators:
        if profile.user and profile.user.email:
            try:
                send_mail(
                    "Message from Brain-O-Math Admin",
                    message_content,
                    settings.DEFAULT_FROM_EMAIL,
                    [profile.user.email],
                    fail_silently=True,
                )
                success_count += 1
            except Exception as e:
                print(f"Error sending email to {profile.user.email}: {e}")
                
    return JsonResponse({"success": True, "data": {"successCount": success_count}})
